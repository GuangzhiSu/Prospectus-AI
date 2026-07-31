"""Document ingest for eligibility — standalone (no ai-module import).

Supports PDF (PyMuPDF), DOCX (python-docx), XLSX (openpyxl/pandas), JSON, and
plain text. Returns page-/sheet-anchored text blocks for the extraction agent.

JSON handling covers:
  - plain text / narrative JSON
  - issuer v3 envelopes
  - Agent1/Agent2 chapter dumps (``values`` + ``extracted_source_materials``)
"""
from __future__ import annotations

import json
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any


@dataclass
class TextBlock:
    """One contiguous text unit with optional page / sheet provenance."""

    text: str
    source_file: str
    page_start: int | None = None
    page_end: int | None = None
    label: str = ""  # e.g. sheet name, "narrative", "table_1"


@dataclass
class DocumentBundle:
    """All text blocks loaded from one or more uploaded files."""

    blocks: list[TextBlock] = field(default_factory=list)
    errors: list[str] = field(default_factory=list)

    def combined_text(self, max_chars: int = 80_000) -> str:
        """Concatenate blocks, preferring eligibility-critical chapters first.

        Oversized blocks are sliced (never skipped entirely) so a large first
        file cannot collapse the prompt to only ``[... truncated ...]``.
        """
        priority_tokens = (
            "summary",
            "financial_information",
            "financial",
            "share_capital",
            "substantial_shareholders",
            "business",
            "history",
            "corporate_information",
            "cornerstone",
            "structure_of_the_global_offering",
            "risk_factors",
        )

        def _rank(block: TextBlock) -> tuple[int, str]:
            name = (block.source_file or "").lower().replace("-", "_").replace(" ", "_")
            for idx, token in enumerate(priority_tokens):
                if token in name:
                    return (idx, name)
            return (len(priority_tokens), name)

        ordered = sorted(self.blocks, key=_rank)
        parts: list[str] = []
        size = 0
        for block in ordered:
            if size >= max_chars:
                parts.append("[... truncated ...]\n")
                break
            header = f"[{block.source_file}"
            if block.label:
                header += f" | {block.label}"
            if block.page_start is not None:
                header += f" | p.{block.page_start}"
                if block.page_end and block.page_end != block.page_start:
                    header += f"-{block.page_end}"
            header += "]\n"
            body = (block.text or "").strip()
            if not body:
                continue
            remaining = max_chars - size - len(header) - 2
            if remaining <= 80:
                parts.append("[... truncated ...]\n")
                break
            if len(body) > remaining:
                body = body[:remaining].rstrip() + "\n[... truncated ...]"
            chunk = header + body + "\n\n"
            parts.append(chunk)
            size += len(chunk)
        return "".join(parts)


def _format_value(value: Any) -> str:
    if value is None:
        return ""
    if isinstance(value, (str, int, float, bool)):
        return str(value)
    return json.dumps(value, ensure_ascii=False)


def _flatten_agent_chapter_json(data: dict[str, Any]) -> str:
    """Turn Agent1 chapter JSON into compact readable text for eligibility LLM."""
    parts: list[str] = []
    notes = (data.get("coverage_notes") or "").strip()
    if notes:
        parts.append(f"Coverage notes: {notes}")

    values = data.get("values")
    if isinstance(values, dict) and values:
        parts.append("Extracted field values:")
        for key, raw in values.items():
            if isinstance(raw, dict) and "value" in raw:
                val = _format_value(raw.get("value"))
                unit = raw.get("unit")
                line = f"- {key}: {val}"
                if unit:
                    line += f" ({unit})"
                preview = (raw.get("span_preview") or "").strip()
                if preview and preview not in val:
                    line += f"\n  context: {preview[:400]}"
                parts.append(line)
            else:
                parts.append(f"- {key}: {_format_value(raw)}")

    mats = data.get("extracted_source_materials")
    if isinstance(mats, dict):
        facts = mats.get("key_numeric_facts") or []
        if isinstance(facts, list) and facts:
            parts.append("Key numeric facts:")
            for item in facts[:40]:
                if isinstance(item, dict):
                    text = (item.get("text") or item.get("value") or "").strip()
                    if text:
                        parts.append(f"- {text[:600]}")
                elif isinstance(item, str) and item.strip():
                    parts.append(f"- {item.strip()[:600]}")

        narrative = mats.get("key_narrative_points") or []
        if isinstance(narrative, list) and narrative:
            parts.append("Key narrative points:")
            for item in narrative[:20]:
                if isinstance(item, dict):
                    text = (item.get("text") or "").strip()
                    if text:
                        parts.append(f"- {text[:600]}")
                elif isinstance(item, str) and item.strip():
                    parts.append(f"- {item.strip()[:600]}")

        excerpts = mats.get("source_excerpt_blocks") or []
        if isinstance(excerpts, list) and excerpts:
            parts.append("Source excerpts:")
            budget = 8_000
            used = 0
            for item in excerpts:
                if not isinstance(item, dict):
                    continue
                text = (item.get("text") or "").strip()
                if not text:
                    continue
                take = text[: min(1_200, budget - used)]
                parts.append(take)
                used += len(take)
                if used >= budget:
                    break

    nulls = data.get("null_reasons")
    if isinstance(nulls, dict) and nulls:
        parts.append("Null reasons:")
        for key, reason in nulls.items():
            parts.append(f"- {key}: {reason}")

    return "\n".join(parts).strip()


def _load_pdf(path: Path) -> list[TextBlock]:
    try:
        import fitz  # PyMuPDF
    except ImportError as exc:
        raise ImportError("PDF support requires pymupdf: pip install pymupdf") from exc
    blocks: list[TextBlock] = []
    with fitz.open(str(path)) as doc:
        for i, page in enumerate(doc, start=1):
            txt = page.get_text("text") or ""
            if not txt.strip():
                continue
            blocks.append(
                TextBlock(
                    text=txt,
                    source_file=path.name,
                    page_start=i,
                    page_end=i,
                    label="page",
                )
            )
    return blocks


def _load_docx(path: Path) -> list[TextBlock]:
    try:
        from docx import Document
    except ImportError as exc:
        raise ImportError(
            "DOCX support requires python-docx: pip install python-docx"
        ) from exc
    doc = Document(str(path))
    blocks: list[TextBlock] = []
    paras = [p.text.strip() for p in doc.paragraphs if p.text and p.text.strip()]
    if paras:
        blocks.append(
            TextBlock(
                text="\n\n".join(paras),
                source_file=path.name,
                label="narrative",
            )
        )
    for idx, tbl in enumerate(doc.tables, start=1):
        rows: list[str] = []
        for row in tbl.rows:
            cells = [c.text.strip().replace("\n", " ") for c in row.cells]
            rows.append(" | ".join(cells))
        if rows:
            blocks.append(
                TextBlock(
                    text=f"[Table {idx}]\n" + "\n".join(rows),
                    source_file=path.name,
                    label=f"table_{idx}",
                )
            )
    return blocks


def _load_xlsx(path: Path) -> list[TextBlock]:
    try:
        import pandas as pd
    except ImportError as exc:
        raise ImportError(
            "XLSX support requires pandas and openpyxl: pip install pandas openpyxl"
        ) from exc
    xl = pd.ExcelFile(path, engine="openpyxl")
    blocks: list[TextBlock] = []
    for sheet in xl.sheet_names:
        df = pd.read_excel(xl, sheet_name=sheet, header=None, dtype=str)
        text = df.to_string(index=False, header=False, na_rep="")
        if text.strip():
            blocks.append(
                TextBlock(
                    text=f"[Sheet: {sheet}]\n{text}",
                    source_file=path.name,
                    label=sheet,
                )
            )
    return blocks


def _load_json(path: Path) -> list[TextBlock]:
    with open(path, encoding="utf-8") as handle:
        data = json.load(handle)

    # Agent1/Agent2 chapter package: prefer flattened values + excerpts.
    if isinstance(data, dict) and (
        isinstance(data.get("values"), dict)
        or isinstance(data.get("extracted_source_materials"), dict)
    ):
        flat = _flatten_agent_chapter_json(data)
        if flat:
            return [
                TextBlock(text=flat, source_file=path.name, label="chapter_json"),
            ]

    # Issuer / generic JSON — dump compactly (no indent) to save context.
    text = json.dumps(data, ensure_ascii=False, separators=(",", ":"))
    return [
        TextBlock(text=text, source_file=path.name, label="json"),
    ]


def _load_text(path: Path) -> list[TextBlock]:
    return [
        TextBlock(
            text=path.read_text(encoding="utf-8", errors="replace"),
            source_file=path.name,
            label="text",
        )
    ]


_LOADERS = {
    ".pdf": _load_pdf,
    ".docx": _load_docx,
    ".xlsx": _load_xlsx,
    ".xls": _load_xlsx,
    ".json": _load_json,
    ".txt": _load_text,
    ".md": _load_text,
}


def load_document(path: str | Path) -> DocumentBundle:
    """Load one document into a ``DocumentBundle``."""
    path = Path(path)
    bundle = DocumentBundle()
    if not path.exists():
        bundle.errors.append(f"file not found: {path}")
        return bundle
    loader = _LOADERS.get(path.suffix.lower())
    if loader is None:
        bundle.errors.append(f"unsupported file type: {path.suffix}")
        return bundle
    try:
        bundle.blocks.extend(loader(path))
    except Exception as exc:  # noqa: BLE001
        bundle.errors.append(f"{path.name}: {exc}")
    return bundle


def load_documents(paths: list[str | Path]) -> DocumentBundle:
    """Load many documents into one combined bundle."""
    combined = DocumentBundle()
    for path in paths:
        part = load_document(path)
        combined.blocks.extend(part.blocks)
        combined.errors.extend(part.errors)
    return combined


def issuer_json_as_bundle(root: dict[str, Any], source_name: str = "issuer.json") -> DocumentBundle:
    """Wrap an already-loaded issuer JSON dict as a document bundle."""
    text = json.dumps(root, ensure_ascii=False, separators=(",", ":"))
    return DocumentBundle(
        blocks=[TextBlock(text=text, source_file=source_name, label="json")]
    )
