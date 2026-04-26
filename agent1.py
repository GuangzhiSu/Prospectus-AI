#!/usr/bin/env python3
"""
Agent1: Process Excel and JSON files → dual retrieval stores.

Output:
  agent1_output/
    text_chunks.jsonl   - Narrative content (Excel, industry descriptions, long text)
    fact_store.jsonl    - Structured facts (metrics, periods, values)
    manifest.json

Text store: chunk_size 500–700, overlap 100. For RAG over narrative.
Fact store: flattened field.period.metric → value, unit, metadata.
"""

from __future__ import annotations

import argparse
import json
import hashlib
import re
from pathlib import Path
from typing import Any

DEFAULT_MODEL = "Qwen/Qwen2.5-7B-Instruct"

# Text store chunk params (narrative generation)
TEXT_CHUNK_SIZE = 600  # 500–700 range
TEXT_CHUNK_OVERLAP = 100

SECTIONS = [
    ("A", "Business & Strategy"),
    ("B", "Industry & Market"),
    ("C", "Risk Factors"),
    ("D", "Financial Performance & Condition"),
    ("E", "Use of Proceeds & Capital Structure"),
    ("F", "Management, Governance & Incentives"),
    ("G", "Legal, Regulatory & Compliance"),
    ("H", "Offering Mechanics & Share Structure"),
]

FILE_TO_SECTION: dict[str, str] = {
    "company-introduction": "A", "business-data": "A", "market-performance-comparison": "B",
    "comprehensive-comparison": "B", "balance-sheet": "D", "financial-ratios-comparison": "D",
    "financial-data-comparison": "D", "cash-flow": "D", "growth-capability": "D",
    "operating-capability": "D", "profit-forecast-comparison": "D",
    "share-capital-structure": "E", "holdings-or-equity": "E", "mainland-fund-holdings": "E",
    "board-and-executives": "F",
    # Reverse-engineered native-doc filenames (slice + template outputs)
    "industry_overview": "B", "regulatory_overview": "G", "legal_opinion": "G",
    "connected_transactions": "G", "vie_opinion": "G",
    "financial_information": "D", "financial_model": "D", "comfort_letter": "D",
    "audit_report": "D", "mdna": "D",
    "sponsor_dd_memorandum": "C", "risk_factors": "C",
    "issuer_corporate_info": "A", "corporate_info": "A",
    "shareholder_list": "E", "share_capital": "E",
    "directors_senior_management": "F",
}

# Filename-token -> Schema A gating-doc field_id (optional; used when manifest absent)
FILENAME_TO_SCHEMA_FIELD: dict[str, str] = {
    "industry_overview": "prof.industry_consultant.research_report",
    "regulatory_overview": "prof.legal.regulatory_opinion",
    "legal_opinion": "prof.legal.hk_legal_opinion",
    "vie_opinion": "prof.legal.vie_opinion",
    "connected_transactions": "prof.legal.connected_transactions_memo",
    "sponsor_dd_memorandum": "prof.sponsor.dd_memorandum",
    "mdna": "prof.accountants.mdna_commentary",
    "financial_model": "prof.accountants.financial_model",
    "audit_report": "prof.accountants.audit_report",
    "comfort_letter": "prof.accountants.comfort_letter",
    "shareholder_list": "issuer.corp.shareholder_register",
    "issuer_corporate_info": "issuer.corp.certificate_of_incorporation",
}

# Canonical prospectus section (Agent2 naming) -> legacy Agent1 letter bucket (A-H)
CANONICAL_TO_LEGACY: dict[str, str] = {
    "Business": "A", "Corporate_Information": "A",
    "History_Reorganization_Corporate_Structure": "A",
    "Industry_Overview": "B",
    "Risk_Factors": "C", "Forward_Looking_Statements": "C",
    "Financial_Information": "D", "Summary": "D",
    "Future_Plans_and_Use_of_Proceeds": "E", "Share_Capital": "E",
    "Substantial_Shareholders": "E", "Cornerstone_Investors": "E",
    "Directors_and_Senior_Management": "F",
    "Relationship_with_Controlling_Shareholders": "F",
    "Regulatory_Overview": "G", "Waivers_and_Exemptions": "G",
    "Connected_Transactions": "G", "Contractual_Arrangements_VIE": "G",
    "Structure_of_the_Global_Offering": "H", "Underwriting": "H",
    "Parties_Involved_in_the_Global_Offering": "H",
    "Expected_Timetable": "H", "How_to_Apply_for_Hong_Kong_Offer_Shares": "H",
}

JSON_CATEGORY_TO_SECTION: dict[str, str] = {
    "company_profile": "A", "corporate_structure": "A", "business": "A",
    "products_and_technology": "A", "customers": "A", "market": "B", "competition": "B",
    "financials": "D", "operating_metrics": "D", "risk_related_data": "C",
    "management": "F", "shareholders": "E", "ipo_offering": "H",
    # legacy keys
    "company_information": "A", "business_operations": "A", "offering_information": "H",
}


def _infer_unit(key: str) -> str | None:
    key_lower = key.lower()
    if "_rmb" in key_lower or "rmb" in key_lower: return "RMB"
    if "_hkd" in key_lower or "hkd" in key_lower: return "HKD"
    if "_usd" in key_lower or "usd" in key_lower: return "USD"
    if "_pct" in key_lower or "_ratio" in key_lower or "pct" in key_lower: return "%"
    if "margin" in key_lower and "gross" in key_lower or "net_margin" in key_lower: return "%"
    return None


def _extract_facts(
    prefix: str,
    value: Any,
    facts: list[dict[str, Any]],
    source_file: str,
) -> None:
    """Recursively extract structured facts from JSON value."""
    if value is None:
        return
    if isinstance(value, (str, int, float, bool)):
        unit = _infer_unit(prefix.split(".")[-1]) if "." in prefix else None
        facts.append({
            "field": prefix,
            "period": None,
            "metric": prefix.split(".")[-1] if "." in prefix else prefix,
            "value": value,
            "unit": unit,
            "metadata": {"source_file": source_file},
        })
        return
    if isinstance(value, list):
        for i, item in enumerate(value):
            if isinstance(item, dict):
                period = item.get("period") or item.get("date")
                for k, v in item.items():
                    if k in ("period", "date", "period_reference"):
                        continue
                    if isinstance(v, (str, int, float, bool)) and v is not None:
                        unit = _infer_unit(k)
                        fact: dict[str, Any] = {
                            "field": prefix,
                            "period": period,
                            "metric": k,
                            "value": v,
                            "unit": unit,
                            "metadata": {"source_file": source_file},
                        }
                        facts.append(fact)
                    elif isinstance(v, (dict, list)):
                        _extract_facts(f"{prefix}.{k}", v, facts, source_file)
            elif isinstance(item, (str, int, float, bool)):
                facts.append({
                    "field": prefix,
                    "period": None,
                    "metric": "item",
                    "value": item,
                    "unit": None,
                    "metadata": {"source_file": source_file, "index": i},
                })
        return
    if isinstance(value, dict):
        # Check for range-style (low, high, currency)
        if "low" in value and "high" in value:
            facts.append({
                "field": prefix,
                "period": None,
                "metric": "price_range",
                "value": value.get("low"),
                "unit": value.get("currency", "HKD"),
                "metadata": {"high": value.get("high"), "source_file": source_file},
            })
            if value.get("high") != value.get("low"):
                facts.append({
                    "field": prefix,
                    "period": None,
                    "metric": "price_range_high",
                    "value": value.get("high"),
                    "unit": value.get("currency", "HKD"),
                    "metadata": {"source_file": source_file},
                })
            return
        for k, v in value.items():
            if isinstance(v, (str, int, float, bool)):
                unit = _infer_unit(k)
                facts.append({
                    "field": f"{prefix}.{k}",
                    "period": None,
                    "metric": k,
                    "value": v,
                    "unit": unit,
                    "metadata": {"source_file": source_file},
                })
            else:
                _extract_facts(f"{prefix}.{k}", v, facts, source_file)


def extract_facts_from_json(path: Path) -> list[dict[str, Any]]:
    """Extract fact entries from structured IPO JSON."""
    with open(path, "r", encoding="utf-8") as f:
        data = json.load(f)
    if not isinstance(data, dict):
        return []
    facts: list[dict[str, Any]] = []
    source = path.name
    for category, fields in data.items():
        if not isinstance(fields, dict):
            continue
        section = JSON_CATEGORY_TO_SECTION.get(category, "A")
        for field_name, value in fields.items():
            if value is None:
                continue
            full_field = f"{category}.{field_name}"
            start_len = len(facts)
            _extract_facts(full_field, value, facts, source)
            for f in facts[start_len:]:
                f.setdefault("metadata", {})["section_hint"] = section
    return facts


def _is_narrative_field(category: str, field_name: str, value: Any) -> bool:
    """Fields suitable for text store: arrays of strings, long text."""
    if isinstance(value, str) and len(value) > 200:
        return True
    if isinstance(value, list) and value and isinstance(value[0], str):
        # industry_trends, competitive_advantages, etc.
        return True
    return False


def _narrative_to_text(prefix: str, value: Any) -> str:
    if isinstance(value, str):
        return value
    if isinstance(value, list) and value and isinstance(value[0], str):
        return "\n".join(f"- {v}" for v in value)
    return ""


def extract_text_from_json(path: Path) -> list[dict[str, Any]]:
    """Extract narrative text chunks from JSON (industry descriptions, long text)."""
    with open(path, "r", encoding="utf-8") as f:
        data = json.load(f)
    if not isinstance(data, dict):
        return []
    texts: list[dict[str, Any]] = []
    source = path.name
    for category, fields in data.items():
        if not isinstance(fields, dict):
            continue
        section = JSON_CATEGORY_TO_SECTION.get(category, "A")
        for field_name, value in fields.items():
            if not _is_narrative_field(category, field_name, value):
                continue
            text = _narrative_to_text(f"{category}.{field_name}", value)
            if not text.strip():
                continue
            texts.append({
                "text": text,
                "source_file": source,
                "section_hint": section,
                "topic": f"{category}.{field_name}",
                "importance": "medium",
            })
    return texts


def chunk_text(text: str, max_chars: int = 600, overlap: int = 100) -> list[str]:
    cleaned = re.sub(r"\r", "\n", text).strip()
    cleaned = re.sub(r"[ \t]+\n", "\n", cleaned)
    cleaned = re.sub(r"\n{3,}", "\n\n", cleaned).strip()
    if not cleaned:
        return []
    chunks: list[str] = []
    i = 0
    while i < len(cleaned):
        end = min(i + max_chars, len(cleaned))
        slice_ = cleaned[i:end].strip()
        if slice_:
            chunks.append(slice_)
        i = end - overlap
        if i < 0:
            i = 0
        if end >= len(cleaned):
            break
    return chunks


def summarize_table_with_qwen(
    text_sample: str,
    filename: str,
    sheet_name: str,
    model_name: str = DEFAULT_MODEL,
) -> str:
    import os

    if os.environ.get("AGENT1_SKIP_TABLE_LLM") == "1":
        return f"Table: {filename} / {sheet_name}"

    prompt = (
        "Summarize this Excel table in 2-4 sentences for HKEX prospectus drafting. "
        "Be factual.\n\n"
        f"File: {filename}\nSheet: {sheet_name}\n\nExcerpt:\n{text_sample[:2500]}\n\nSUMMARY:"
    )
    try:
        provider = (os.environ.get("LLM_PROVIDER") or "qwen_local").strip().lower()
        if provider == "openai":
            from llm_openai import run_openai_chat

            out = run_openai_chat(prompt, max_new_tokens=256)
            return out.strip() or f"Table: {filename} / {sheet_name}"
        from llm_qwen import run_qwen_text

        out = run_qwen_text(
            prompt,
            model_name=model_name,
            max_new_tokens=256,
        )
        return out.strip() or f"Table: {filename} / {sheet_name}"
    except Exception:
        return f"Table: {filename} / {sheet_name}"


def classify_file(filename: str) -> str:
    stem = Path(filename).stem.lower()
    for key, section_id in FILE_TO_SECTION.items():
        if key in stem:
            return section_id
    return "D"


def classify_file_with_schema(
    filename: str,
    manifest_lookup: dict[str, dict] | None = None,
) -> tuple[str, str | None, str | None]:
    """Resolve (legacy_letter_section, schema_a_field_id, canonical_prospectus_section).

    Priority: native_docs manifest entry > filename-token map > legacy keyword map.
    `canonical_prospectus_section` uses Agent2 section naming (e.g. ``Industry_Overview``);
    ``None`` when unknown.
    """
    canonical: str | None = None
    schema_field: str | None = None
    if manifest_lookup:
        entry = (
            manifest_lookup.get(filename)
            or manifest_lookup.get(Path(filename).name)
        )
        if entry:
            schema_field = entry.get("schema_a_field_id") or None
            canonical = entry.get("section_hint") or None
            legacy = CANONICAL_TO_LEGACY.get(canonical or "") if canonical else None
            if not legacy:
                legacy = classify_file(filename)
            return legacy, schema_field, canonical

    stem = Path(filename).stem.lower()
    for token, field_id in FILENAME_TO_SCHEMA_FIELD.items():
        if token in stem:
            schema_field = field_id
            break
    legacy = classify_file(filename)
    return legacy, schema_field, canonical


def _load_input_manifest(data_path: Path) -> dict[str, dict]:
    """Load ``<data_dir>/manifest.json`` (from native-doc reverse-engineering).

    Returns a ``{path_or_basename: entry}`` lookup. Silent no-op when absent.
    """
    mpath = data_path / "manifest.json"
    if not mpath.exists():
        return {}
    try:
        raw = json.loads(mpath.read_text(encoding="utf-8"))
    except Exception:  # noqa: BLE001
        return {}
    lookup: dict[str, dict] = {}
    for entry in raw.get("files", []) or []:
        p = entry.get("path", "")
        if not p:
            continue
        lookup[p] = entry
        lookup[Path(p).name] = entry
    return lookup


def extract_from_docx(path: Path) -> list[tuple[str, str]]:
    """Return ``(topic, text)`` tuples for a .docx: one for narrative, one per table."""
    try:
        from docx import Document
    except ImportError as exc:
        raise ImportError(
            "agent1 DOCX support requires python-docx: pip install python-docx"
        ) from exc
    doc = Document(str(path))
    out: list[tuple[str, str]] = []
    paras = [p.text.strip() for p in doc.paragraphs if p.text and p.text.strip()]
    if paras:
        out.append(("narrative", "\n\n".join(paras)))
    for idx, tbl in enumerate(doc.tables, start=1):
        rows: list[str] = []
        for row in tbl.rows:
            cells = [c.text.strip().replace("\n", " ") for c in row.cells]
            rows.append(" | ".join(cells))
        if rows:
            out.append((f"table_{idx}", f"[Table {idx}]\n" + "\n".join(rows)))
    return out


def extract_from_pdf(path: Path) -> list[tuple[int, str]]:
    """Return per-page ``(page_number_1indexed, text)`` for a PDF (no OCR)."""
    try:
        import fitz  # PyMuPDF
    except ImportError as exc:
        raise ImportError(
            "agent1 PDF support requires PyMuPDF: pip install pymupdf"
        ) from exc
    out: list[tuple[int, str]] = []
    with fitz.open(str(path)) as doc:
        for i, page in enumerate(doc, start=1):
            txt = page.get_text("text") or ""
            if len(txt.strip()) < 20:
                print(
                    f"  [PDF warn] {path.name} page {i}: <20 chars "
                    "(likely image-based; OCR deferred)"
                )
                if not txt.strip():
                    continue
            out.append((i, txt))
    return out


def extract_per_sheet(path: Path) -> list[tuple[str, str]]:
    try:
        import pandas as pd
    except ImportError:
        raise ImportError("agent1 requires pandas and openpyxl")
    xl = pd.ExcelFile(path, engine="openpyxl")
    result: list[tuple[str, str]] = []
    for sheet in xl.sheet_names:
        df = pd.read_excel(xl, sheet_name=sheet, header=None, dtype=str)
        text = df.to_string(index=False, header=False, na_rep="")
        if text.strip():
            result.append((sheet, f"[Sheet: {sheet}]\n{text}"))
    return result


def run_agent1(
    data_dir: str | Path = "data",
    output_dir: str | Path = "agent1_output",
    text_chunk_size: int = TEXT_CHUNK_SIZE,
    text_chunk_overlap: int = TEXT_CHUNK_OVERLAP,
    model_name: str = DEFAULT_MODEL,
    project_id: str | None = None,
) -> None:
    # Per-project isolation: when project_id set, resolve data_dir + output_dir
    # under that namespace unless the caller already points at a project folder.
    base_data = Path(data_dir)
    base_out = Path(output_dir)
    if project_id:
        candidate = base_data / project_id
        data_path = candidate if candidate.exists() else base_data
        output_path = base_out / project_id
    else:
        data_path = base_data
        output_path = base_out
    output_path.mkdir(parents=True, exist_ok=True)
    (output_path / "by_section").mkdir(exist_ok=True)

    manifest_lookup = _load_input_manifest(data_path)
    if manifest_lookup:
        print(
            f"  [manifest] loaded {len(manifest_lookup)//2} entries from "
            f"{data_path/'manifest.json'}"
        )

    excel_files = sorted(data_path.rglob("*.xlsx"))
    json_files = sorted(data_path.glob("*.json"))
    # Exclude the input manifest.json from JSON-ingest path.
    json_files = [p for p in json_files if p.name != "manifest.json"]
    docx_files = sorted(data_path.rglob("*.docx"))
    pdf_files = sorted(data_path.rglob("*.pdf"))
    if not any([excel_files, json_files, docx_files, pdf_files]):
        raise FileNotFoundError(
            f"No .xlsx / .json / .docx / .pdf files in {data_path}. "
            "Put files there first."
        )

    text_chunks: list[dict[str, Any]] = []
    fact_store: list[dict[str, Any]] = []
    sheet_summaries: dict[tuple[str, str], str] = {}
    by_section: dict[str, list[dict[str, Any]]] = {s[0]: [] for s in SECTIONS}

    # --- Excel → text store ---
    for f in excel_files:
        sheets = extract_per_sheet(f)
        if not sheets:
            continue
        section_id, schema_field, canonical = classify_file_with_schema(
            f.name, manifest_lookup
        )
        section_name = next((n for sid, n in SECTIONS if sid == section_id), "Unknown")
        for sheet_name, sheet_text in sheets:
            print(f"  [Qwen] {f.name} / {sheet_name} -> summarising...")
            summary = summarize_table_with_qwen(sheet_text, f.name, sheet_name, model_name=model_name)
            sheet_summaries[(f.name, sheet_name)] = summary
            chunks = chunk_text(sheet_text, max_chars=text_chunk_size, overlap=text_chunk_overlap)
            for i, chunk in enumerate(chunks):
                chunk_id = hashlib.md5(f"{f.name}:{sheet_name}:{i}:{chunk[:50]}".encode()).hexdigest()[:12]
                rec = {
                    "text": f"{summary}\n\n[Data]\n{chunk}",
                    "source_file": f.name,
                    "section_hint": section_id,
                    "schema_field_hint": schema_field,
                    "prospectus_section_hint": canonical,
                    "topic": sheet_name,
                    "importance": "high",
                    "chunk_id": chunk_id,
                    "source_type": "excel_sheet",
                    "sheet_name": sheet_name,
                }
                text_chunks.append(rec)
                by_section[section_id].append(rec)
        print(f"  [Qwen] {f.name} -> Section {section_id} ({len(sheets)} sheets)")

    # --- JSON → fact store + narrative text ---
    for f in json_files:
        print(f"  [JSON] {f.name} -> extracting...")
        facts = extract_facts_from_json(f)
        for i, fact in enumerate(facts):
            fact["fact_id"] = hashlib.md5(
                f"{f.name}:{fact.get('field','')}:{i}:{str(fact.get('value',''))}".encode()
            ).hexdigest()[:12]
            fact_store.append(fact)
        narrative = extract_text_from_json(f)
        for item in narrative:
            chunk_list = chunk_text(
                item["text"],
                max_chars=text_chunk_size,
                overlap=text_chunk_overlap,
            )
            for j, chunk in enumerate(chunk_list):
                chunk_id = hashlib.md5(f"{f.name}:{item['topic']}:{j}".encode()).hexdigest()[:12]
                rec = {
                    "text": chunk,
                    "source_file": item["source_file"],
                    "section_hint": item["section_hint"],
                    "topic": item["topic"],
                    "importance": item["importance"],
                    "chunk_id": chunk_id,
                    "source_type": "json_narrative",
                }
                text_chunks.append(rec)
                by_section[item["section_hint"]].append(rec)
        print(f"  [JSON] {f.name} -> {len(facts)} facts, {len(narrative)} narrative blocks")

    # --- DOCX → text store ---
    for f in docx_files:
        try:
            blocks = extract_from_docx(f)
        except ImportError as exc:
            print(f"  [DOCX skip] {f.name}: {exc}")
            break
        except Exception as exc:  # noqa: BLE001
            print(f"  [DOCX error] {f.name}: {exc}")
            continue
        if not blocks:
            continue
        section_id, schema_field, canonical = classify_file_with_schema(
            f.name, manifest_lookup
        )
        for topic, text in blocks:
            chunks = chunk_text(
                text, max_chars=text_chunk_size, overlap=text_chunk_overlap
            )
            for i, chunk in enumerate(chunks):
                chunk_id = hashlib.md5(
                    f"{f.name}:{topic}:{i}:{chunk[:50]}".encode()
                ).hexdigest()[:12]
                rec = {
                    "text": chunk,
                    "source_file": f.name,
                    "section_hint": section_id,
                    "schema_field_hint": schema_field,
                    "prospectus_section_hint": canonical,
                    "topic": topic,
                    "importance": "high" if topic.startswith("table_") else "medium",
                    "chunk_id": chunk_id,
                    "source_type": "docx_table" if topic.startswith("table_") else "docx_paragraphs",
                }
                text_chunks.append(rec)
                by_section[section_id].append(rec)
        print(
            f"  [DOCX] {f.name} -> Section {section_id}"
            f"{' / '+schema_field if schema_field else ''} ({len(blocks)} block(s))"
        )

    # --- Raw PDF → text store ---
    for f in pdf_files:
        try:
            pages = extract_from_pdf(f)
        except ImportError as exc:
            print(f"  [PDF skip] {f.name}: {exc}")
            break
        except Exception as exc:  # noqa: BLE001
            print(f"  [PDF error] {f.name}: {exc}")
            continue
        if not pages:
            continue
        section_id, schema_field, canonical = classify_file_with_schema(
            f.name, manifest_lookup
        )
        for page_no, page_text in pages:
            chunks = chunk_text(
                page_text, max_chars=text_chunk_size, overlap=text_chunk_overlap
            )
            for i, chunk in enumerate(chunks):
                chunk_id = hashlib.md5(
                    f"{f.name}:p{page_no}:{i}:{chunk[:50]}".encode()
                ).hexdigest()[:12]
                rec = {
                    "text": chunk,
                    "source_file": f.name,
                    "section_hint": section_id,
                    "schema_field_hint": schema_field,
                    "prospectus_section_hint": canonical,
                    "topic": f"page_{page_no}",
                    "importance": "medium",
                    "chunk_id": chunk_id,
                    "source_type": "pdf_page",
                    "page": page_no,
                }
                text_chunks.append(rec)
                by_section[section_id].append(rec)
        print(
            f"  [PDF] {f.name} -> Section {section_id}"
            f"{' / '+schema_field if schema_field else ''} ({len(pages)} page(s))"
        )

    # --- Write text_chunks.jsonl ---
    text_path = output_path / "text_chunks.jsonl"
    with open(text_path, "w", encoding="utf-8") as out:
        for rec in text_chunks:
            out.write(json.dumps(rec, ensure_ascii=False) + "\n")
    print(f"Wrote {len(text_chunks)} text chunks to {text_path}")

    # --- Write fact_store.jsonl ---
    fact_path = output_path / "fact_store.jsonl"
    with open(fact_path, "w", encoding="utf-8") as out:
        for rec in fact_store:
            out.write(json.dumps(rec, ensure_ascii=False) + "\n")
    print(f"Wrote {len(fact_store)} facts to {fact_path}")

    # --- Backward compatibility: rag_chunks.jsonl (merged view for old retriever) ---
    rag_chunks: list[dict[str, Any]] = []
    for i, rec in enumerate(text_chunks):
        rag_chunks.append({
            "chunk_id": rec.get("chunk_id", str(i)),
            "section_id": rec.get("section_hint", "A"),
            "section": next((n for sid, n in SECTIONS if sid == rec.get("section_hint", "A")), "Unknown"),
            "source_file": rec.get("source_file", ""),
            "sheet_name": rec.get("topic", rec.get("sheet_name", "")),
            "chunk_index": i,
            "text": rec["text"],
            "sheet_summary": rec.get("topic", "")[:100],
            "source_type": rec.get("source_type", "text_chunk"),
        })
    # Append fact summaries as text-like chunks for context (so section writer sees facts)
    facts_by_section: dict[str, list[dict]] = {}
    for fact in fact_store:
        sh = fact.get("metadata", {}).get("section_hint", "A")
        facts_by_section.setdefault(sh, []).append(fact)
    for section_id, facts in facts_by_section.items():
        block = "\n".join(
            f"{f['field']}: {f['metric']}={f['value']}" + (f" ({f['period']})" if f.get("period") else "")
            for f in facts[:50]  # limit to avoid huge context
        )
        if block:
            rag_chunks.append({
                "chunk_id": hashlib.md5(f"facts_{section_id}".encode()).hexdigest()[:12],
                "section_id": section_id,
                "section": next((n for sid, n in SECTIONS if sid == section_id), "Unknown"),
                "source_file": "fact_store",
                "sheet_name": f"facts_{section_id}",
                "chunk_index": len(rag_chunks),
                "text": f"[Structured Facts - {section_id}]\n{block}",
                "sheet_summary": f"Facts for section {section_id}",
                "source_type": "fact_store",
            })
    main_path = output_path / "rag_chunks.jsonl"
    with open(main_path, "w", encoding="utf-8") as out:
        for rec in rag_chunks:
            out.write(json.dumps(rec, ensure_ascii=False) + "\n")
    print(f"Wrote {len(rag_chunks)} combined chunks (backward compat) to {main_path}")

    # --- Write by_section (for section preview API) ---
    for section_id, section_name in SECTIONS:
        recs = by_section[section_id]
        if recs:
            spath = output_path / "by_section" / f"section_{section_id}.jsonl"
            with open(spath, "w", encoding="utf-8") as out:
                for r in recs:
                    # API expects source_file, sheet_name, sheet_summary
                    api_rec = {
                        "source_file": r.get("source_file", ""),
                        "sheet_name": r.get("sheet_name", r.get("topic", "")),
                        "sheet_summary": r.get("topic", r.get("sheet_name", ""))[:150],
                        "text": r.get("text", "")[:500],
                    }
                    out.write(json.dumps(api_rec, ensure_ascii=False) + "\n")
            print(f"  Section {section_id}: {len(recs)} items -> {spath}")

    # --- Write manifest ---
    manifest = {
        "text_chunk_count": len(text_chunks),
        "fact_count": len(fact_store),
        "total_chunks": len(rag_chunks),  # backward compat for frontend
        "sections": [
            {
                "id": s[0],
                "name": s[1],
                "chunk_count": len(by_section[s[0]]),
                "fact_count": len(facts_by_section.get(s[0], [])),
            }
            for s in SECTIONS
        ],
        "source_files": list({c.get("source_file", "") for c in text_chunks}),
        "sheet_summaries": {f"{k[0]}:{k[1]}": v for k, v in sheet_summaries.items()},
    }
    with open(output_path / "manifest.json", "w", encoding="utf-8") as f:
        json.dump(manifest, f, indent=2, ensure_ascii=False)
    print(f"Wrote manifest to {output_path / 'manifest.json'}")


def main() -> None:
    parser = argparse.ArgumentParser(
        description=(
            "Agent1: Excel + JSON + DOCX + PDF -> text_chunks + fact_store. "
            "When --project-id is given, data is read from <data-dir>/<project-id>/ "
            "(if present) and output lands under <output-dir>/<project-id>/. "
            "When <data-dir> contains a manifest.json from reverse_engineer_all.py, "
            "per-file section_hint and Schema A field_id are taken from the manifest."
        )
    )
    parser.add_argument("--data-dir", default="data")
    parser.add_argument("--output-dir", default="agent1_output")
    parser.add_argument(
        "--project-id",
        default=None,
        help="Optional project namespace for multi-run isolation.",
    )
    parser.add_argument("--text-chunk-size", type=int, default=TEXT_CHUNK_SIZE)
    parser.add_argument("--text-chunk-overlap", type=int, default=TEXT_CHUNK_OVERLAP)
    parser.add_argument("--model", default=DEFAULT_MODEL)
    args = parser.parse_args()
    run_agent1(
        data_dir=args.data_dir,
        output_dir=args.output_dir,
        text_chunk_size=args.text_chunk_size,
        text_chunk_overlap=args.text_chunk_overlap,
        model_name=args.model,
        project_id=args.project_id,
    )


if __name__ == "__main__":
    main()
