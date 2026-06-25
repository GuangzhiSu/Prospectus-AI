#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
Patent submission: export curated project source code to a Word document.

Layout requirements:
  - A4 paper
  - Exactly 50 lines per page
  - 60 pages total = 3,000 lines
  - Consolas font, 7.5 pt, line spacing exactly 9 pt
  - 1.2 cm margins on all sides
  - Global line numbers (0001 … 3000)
  - Page header: software name + version

Files are collected in PRIORITY_FILES order (most IP-representative first)
and the budget is consumed greedily — each file is included in full until the
3,000-line limit is reached; the last file in budget is truncated at the exact
line that fills the document.

Usage (run from the project root):
  python export_code_to_word.py
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

# ---------------------------------------------------------------------------
# Document settings
# ---------------------------------------------------------------------------

PROJECT_ROOT = Path(__file__).parent.resolve()
OUTPUT_DOCX = PROJECT_ROOT / "Prospectus_AI_Programme_Code.docx"

SOFTWARE_NAME = "Prospectus AI"
VERSION = "V0.1.0"
TOTAL_PAGES = 60
LINES_PER_PAGE = 50
BUDGET = TOTAL_PAGES * LINES_PER_PAGE  # 3 000

# ---------------------------------------------------------------------------
# Curated file list — ordered from most to least IP-representative.
# Paths are relative to PROJECT_ROOT.
# The collector stops as soon as the 3,000-line budget is consumed.
# ---------------------------------------------------------------------------

PRIORITY_FILES: list[str] = [
    # ── Core Python agents ────────────────────────────────────────────────
    "ai-module/agent2.py",
    "ai-module/agent1.py",
    "ai-module/llm_qwen.py",
    "ai-module/llm_openai.py",
    # ── Knowledge-graph & retrieval layer ────────────────────────────────
    "ai-module/prospectus_graph/retrievers.py",
    "ai-module/prospectus_graph/verifier.py",
    "ai-module/prospectus_graph/output_bundle.py",
    "ai-module/prospectus_graph/crosswalk.py",
    "ai-module/prospectus_graph/blocker_gate.py",
    "ai-module/prospectus_graph/timetable_template.py",
    "ai-module/prospectus_graph/config.py",
    "ai-module/prospectus_graph/issuer_metadata.py",
    "ai-module/prospectus_graph/ai_tag_schema.py",
    "ai-module/prospectus_graph/state.py",
    "ai-module/prospectus_graph/graph.py",
    "ai-module/prospectus_graph/locked_snippets.py",
    # ── Document-graph pipeline ───────────────────────────────────────────
    "knowledge-module/prospectus_docgraph/ingestion/ingestor.py",
    "knowledge-module/prospectus_docgraph/normalizer/title_normalizer.py",
    "knowledge-module/prospectus_docgraph/mining/corpus_miner.py",
    "knowledge-module/prospectus_docgraph/mining/collector.py",
    "knowledge-module/prospectus_docgraph/graph/manager.py",
    "knowledge-module/prospectus_docgraph/export/bundle.py",
    "knowledge-module/prospectus_docgraph/models/nodes.py",
    "knowledge-module/prospectus_docgraph/export/generator_examples.py",
    "knowledge-module/prospectus_docgraph/mining/report_models.py",
    "knowledge-module/prospectus_docgraph/mining/refinement.py",
    "knowledge-module/prospectus_docgraph/export/planner_examples.py",
    "knowledge-module/prospectus_docgraph/export/schema_cards.py",
    "knowledge-module/prospectus_docgraph/export/models.py",
    "knowledge-module/prospectus_docgraph/parser/structure.py",
    "knowledge-module/prospectus_docgraph/schema/loaders.py",
    "knowledge-module/prospectus_docgraph/export/cli.py",
    "knowledge-module/prospectus_docgraph/schema/seed.py",
    "knowledge-module/prospectus_docgraph/export/alias_dataset.py",
    "knowledge-module/prospectus_docgraph/export/serializers.py",
    "knowledge-module/prospectus_docgraph/export/loaders.py",
    "knowledge-module/prospectus_docgraph/mining/models.py",
    "knowledge-module/prospectus_docgraph/models/enums.py",
    "knowledge-module/prospectus_docgraph/normalizer/match_result.py",
    "knowledge-module/prospectus_docgraph/mining/config.py",
    "knowledge-module/prospectus_docgraph/models/edges.py",
    "knowledge-module/prospectus_docgraph/normalizer/aliases_seed.py",
    "knowledge-module/prospectus_docgraph/export/__init__.py",
    # ── Next.js web application (frontend/web is the runnable app) ───────────
    "frontend/web/src/lib/rag.ts",
    "frontend/web/src/lib/app-settings.ts",
    "frontend/web/src/lib/draft-cleaning.ts",
    "frontend/web/src/lib/draft-markdown.ts",
    "frontend/web/src/lib/rag-env.ts",
    "frontend/web/src/lib/prospectus-root.ts",
    "frontend/web/src/app/api/chat/route.ts",
    "frontend/web/src/app/api/agent2/run/route.ts",
    "frontend/web/src/app/api/agent2/export/docx/route.ts",
    "frontend/web/src/app/api/agent2/status/route.ts",
    "frontend/web/src/app/api/agent2/clear/route.ts",
    "frontend/web/src/app/api/agent2/draft/route.ts",
    "frontend/web/src/app/api/agent1/run/route.ts",
    "frontend/web/src/app/api/agent1/upload/route.ts",
    "frontend/web/src/app/api/agent1/section/[id]/route.ts",
    "frontend/web/src/app/api/agent1/results/route.ts",
    "frontend/web/src/app/api/agent1/files/route.ts",
    "frontend/web/src/app/api/agent1/clear-data/route.ts",
    "frontend/web/src/app/api/reset/route.ts",
    "frontend/web/src/app/api/progress/route.ts",
    "frontend/web/src/app/api/files/route.ts",
    "frontend/web/src/app/api/settings/route.ts",
    "frontend/web/src/app/api/settings/test-openai/route.ts",
    "frontend/web/src/app/api/system/gpu/route.ts",
    "frontend/web/src/app/api/models/download/route.ts",
    "frontend/web/src/app/api/models/status/route.ts",
    "frontend/web/src/components/SectionMarkdown.tsx",
    "frontend/web/src/components/AppBackendStatus.tsx",
    "frontend/web/src/app/page.tsx",
    "frontend/web/src/app/agent1/page.tsx",
    "frontend/web/src/app/settings/page.tsx",
    "frontend/web/src/app/kg-view/page.tsx",
    "frontend/web/src/app/layout.tsx",
    "frontend/web/next.config.ts",
    # ── Local LLM service ─────────────────────────────────────────────────
    "platform/services/local-llm/app.py",
]

# Strings that indicate a line holds a secret value (case-insensitive match).
_SENSITIVE_KEYWORDS = {
    "OPENAI_API_KEY",
    "DASHSCOPE_API_KEY",
    "QWEN_API_KEY",
    "API_KEY",
    "SECRET",
    "TOKEN",
    "PASSWORD",
}


# ---------------------------------------------------------------------------
# Line collection
# ---------------------------------------------------------------------------

def _safe_read(path: Path) -> list[str]:
    for enc in ("utf-8", "utf-8-sig", "latin-1"):
        try:
            return path.read_text(encoding=enc, errors="strict").splitlines()
        except Exception:
            continue
    return path.read_text(encoding="utf-8", errors="ignore").splitlines()


def _redact(line: str) -> str:
    cleaned = line.replace("\t", "    ")
    upper = cleaned.upper()
    if any(kw in upper for kw in _SENSITIVE_KEYWORDS):
        # Keep the variable name but hide the value
        if "=" in cleaned:
            left = cleaned[: cleaned.index("=") + 1]
            return f'{left} "<REDACTED>"'
        return "# <REDACTED SENSITIVE CONFIGURATION>"
    return cleaned


def collect_lines() -> list[str]:
    """
    Walk PRIORITY_FILES in order, appending each file's lines to the output
    buffer.  Stop as soon as the budget (3 000 lines) would be exceeded,
    truncating the last file if necessary.

    Each file is preceded by a single separator comment and followed by one
    blank line so the document is easy to navigate.
    """
    output: list[str] = []
    remaining = BUDGET

    for rel in PRIORITY_FILES:
        if remaining <= 0:
            break

        path = PROJECT_ROOT / rel
        if not path.is_file():
            continue

        raw_lines = _safe_read(path)
        redacted = [_redact(ln) for ln in raw_lines]

        # separator + blank line at the end cost 2 lines of budget
        overhead = 2
        available_for_code = remaining - overhead
        if available_for_code <= 0:
            break

        separator = f"# {'─' * 68}"
        file_label = f"# File: {rel}"

        if len(redacted) <= available_for_code - 1:
            # -1 because file_label is the second separator line
            block = [separator, file_label] + redacted + [""]
        else:
            # Truncate this file to fit exactly
            take = available_for_code - 1  # -1 for file_label line
            block = [separator, file_label] + redacted[:take] + [""]

        output.extend(block)
        remaining -= len(block)

    # Pad to exactly BUDGET lines with blank lines if somehow short
    while len(output) < BUDGET:
        output.append("")

    return output[:BUDGET]


# ---------------------------------------------------------------------------
# Word document generation
# ---------------------------------------------------------------------------

def create_docx(lines: list[str]) -> None:
    doc = Document()

    # ── Page layout ───────────────────────────────────────────────────────
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.2)
    section.right_margin = Cm(1.2)

    # ── Header ────────────────────────────────────────────────────────────
    hdr_para = section.header.paragraphs[0]
    hdr_para.text = f"{SOFTWARE_NAME} {VERSION}  –  Programme Code"
    hdr_para.style = doc.styles["Header"]
    for run in hdr_para.runs:
        run.font.name = "Consolas"
        run.font.size = Pt(7)

    # ── Body style ────────────────────────────────────────────────────────
    normal = doc.styles["Normal"]
    normal.font.name = "Consolas"
    normal.font.size = Pt(7.5)
    # Ensure CJK font fallback also uses Consolas
    try:
        normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
    except Exception:
        pass

    # ── Write every line ─────────────────────────────────────────────────
    total = len(lines)
    for i, line in enumerate(lines, start=1):
        numbered = f"{i:04d}  {line}"

        para = doc.add_paragraph()
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(0)
        # "Exactly" line spacing keeps 50 lines firmly on each page
        from docx.oxml import OxmlElement
        pPr = para._p.get_or_add_pPr()
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:line"), "108")   # 108 twentieths-of-a-point = 7.5 pt exactly
        spacing.set(qn("w:lineRule"), "exact")
        spacing.set(qn("w:before"), "0")
        spacing.set(qn("w:after"), "0")
        pPr.append(spacing)

        run = para.add_run(numbered)
        run.font.name = "Consolas"
        run.font.size = Pt(7.5)

        # Insert page break after every 50th line (except the very last)
        if i % LINES_PER_PAGE == 0 and i < total:
            run.add_break(WD_BREAK.PAGE)

    doc.save(str(OUTPUT_DOCX))


# ---------------------------------------------------------------------------
# Entry point
# ---------------------------------------------------------------------------

def main() -> None:
    print(f"Project root : {PROJECT_ROOT}")
    print(f"Budget       : {BUDGET} lines  ({TOTAL_PAGES} pages × {LINES_PER_PAGE} lines/page)")

    lines = collect_lines()
    print(f"Lines collected: {len(lines)}")

    create_docx(lines)

    pages_written = (len(lines) + LINES_PER_PAGE - 1) // LINES_PER_PAGE
    print(f"Pages written  : {pages_written}")
    print(f"Output         : {OUTPUT_DOCX}")


if __name__ == "__main__":
    main()
