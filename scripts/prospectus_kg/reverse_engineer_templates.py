#!/usr/bin/env python3
"""Reverse-engineer source/raw DD deliverables by template filling.

Reads ``prospectus_kg_output/inputs/records/<doc_id>.json`` (Schema B per-section
records) and authors four phase-1 native DD deliverables per document:

    * issuer_corporate_info.docx
    * shareholder_list.xlsx
    * financial_model.xlsx
    * comfort_letter_draft.docx

Outputs land under ``prospectus_kg_output/native_docs/<doc_id>/templated/``.

Any field whose record value is ``null`` is emitted as the literal string
``**DATA_MISSING**`` and surfaced in the returned ``FillResult.missing_fields``
list so the driver can aggregate it in the manifest. No fabrication.

See ``templates/native_docs/README.md`` for the field mapping.
"""

from __future__ import annotations

import argparse
import json
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Iterable

DATA_MISSING = "**DATA_MISSING**"


@dataclass
class FillResult:
    path: Path
    relative_path: str
    schema_a_field_id: str
    section_hint: str
    source_record_paths: list[str]
    missing_fields: list[str] = field(default_factory=list)


# ------------------------------------------------------------------------- utils


def _resolve(rec: dict[str, Any], section_key: str, dotted: str) -> Any:
    """Return the raw stored value for ``<section_key>`` / ``<dotted>``.

    Values may be either:
      - a scalar string/number/None (legacy shape), or
      - an evidence-pointer dict ``{value, source_file, page_start, ...}``.
    """
    sec = rec.get(section_key) or {}
    return sec.get(dotted)


def _text_of(raw: Any) -> str | None:
    """Extract a renderable string from either record shape. Returns ``None``
    when the field is truly absent / null."""
    if raw is None:
        return None
    if isinstance(raw, dict):
        v = raw.get("value")
        if v is None:
            return None
        return str(v)
    return str(raw)


def _pull(
    rec: dict[str, Any],
    section_key: str,
    dotted: str,
    *,
    missing: list[str],
) -> str:
    text = _text_of(_resolve(rec, section_key, dotted))
    if text is None or not text.strip():
        missing.append(f"{section_key}.{dotted}")
        return DATA_MISSING
    return text


# -------------------------------------------------------------- issuer_corporate_info


def _fill_issuer_corporate_info(
    rec: dict[str, Any], out_dir: Path, doc_id: str
) -> FillResult:
    from docx import Document
    from docx.shared import Pt

    missing: list[str] = []
    sec = "section_Corporate_Information"
    cover = "section_Cover"
    rows: list[tuple[str, str]] = [
        ("Issuer name (EN)", _pull(rec, cover, "Cover.issuer_name_en", missing=missing)),
        ("Issuer name (ZH)", _pull(rec, cover, "Cover.issuer_name_ch", missing=missing)),
        ("Stock code", _pull(rec, cover, "Cover.stock_code", missing=missing)),
        (
            "Place / form of incorporation",
            _pull(rec, cover, "Cover.corporate_structure", missing=missing),
        ),
        (
            "Registered office",
            _pull(rec, sec, "Corporate_Information.registered_office_address", missing=missing),
        ),
        (
            "Head office",
            _pull(rec, sec, "Corporate_Information.head_office_address", missing=missing),
        ),
        (
            "Principal place of business (HK)",
            _pull(
                rec, sec, "Corporate_Information.principal_place_of_business_hk", missing=missing
            ),
        ),
        (
            "Company website",
            _pull(rec, sec, "Corporate_Information.company_website", missing=missing),
        ),
        (
            "Joint company secretaries",
            _pull(
                rec, sec, "Corporate_Information.joint_company_secretaries", missing=missing
            ),
        ),
        (
            "Authorized representatives",
            _pull(
                rec, sec, "Corporate_Information.authorized_representatives", missing=missing
            ),
        ),
        (
            "Compliance advisor",
            _pull(rec, sec, "Corporate_Information.compliance_advisor", missing=missing),
        ),
        (
            "Share registrar",
            _pull(rec, sec, "Corporate_Information.share_registrars", missing=missing),
        ),
        (
            "Principal banks",
            _pull(rec, sec, "Corporate_Information.principal_banks", missing=missing),
        ),
        (
            "Audit committee",
            _pull(
                rec, sec, "Corporate_Information.audit_committee_members", missing=missing
            ),
        ),
        (
            "Remuneration committee",
            _pull(
                rec,
                sec,
                "Corporate_Information.remuneration_and_evaluation_committee_members",
                missing=missing,
            ),
        ),
        (
            "Nomination committee",
            _pull(
                rec, sec, "Corporate_Information.nomination_committee_members", missing=missing
            ),
        ),
    ]

    doc = Document()
    heading = doc.add_heading("Issuer Corporate Information", level=0)
    heading.runs[0].font.size = Pt(18)
    doc.add_paragraph(
        f"Document ID: {doc_id}. Reverse-engineered from historical prospectus records; "
        "fields marked **DATA_MISSING** require counsel follow-up."
    )
    tbl = doc.add_table(rows=1 + len(rows), cols=2)
    tbl.style = "Light Grid Accent 1"
    hdr = tbl.rows[0].cells
    hdr[0].text = "Field"
    hdr[1].text = "Value"
    for i, (k, v) in enumerate(rows, start=1):
        tbl.rows[i].cells[0].text = k
        tbl.rows[i].cells[1].text = v

    out_path = out_dir / "issuer_corporate_info.docx"
    out_dir.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    return FillResult(
        path=out_path,
        relative_path="templated/issuer_corporate_info.docx",
        schema_a_field_id="issuer.corp.certificate_of_incorporation",
        section_hint="Corporate_Information",
        source_record_paths=[f"{sec}.*", f"{cover}.*"],
        missing_fields=missing,
    )


# ------------------------------------------------------------------- shareholder_list


def _fill_shareholder_list(
    rec: dict[str, Any], out_dir: Path, doc_id: str
) -> FillResult:
    from openpyxl import Workbook

    missing: list[str] = []
    sec = "section_Substantial_Shareholders"
    cap_sec = "section_Share_Capital"

    wb = Workbook()
    ws = wb.active
    ws.title = "Shareholders"
    ws.append(["Source document:", f"{doc_id}.pdf"])
    ws.append(
        [
            "Note:",
            "Fields emitted as DATA_MISSING must be populated from the issuer's "
            "shareholder register before reliance.",
        ]
    )
    ws.append([])
    ws.append(
        [
            "Shareholder name",
            "Capacity / nature of interest",
            "Number and class of shares",
            "% of class",
            "% of issued share capital",
        ]
    )
    ws.append(
        [
            _pull(rec, sec, "Substantial_Shareholders.name_of_substantial_shareholder", missing=missing),
            _pull(rec, sec, "Substantial_Shareholders.capacity_nature_of_interest", missing=missing),
            _pull(rec, sec, "Substantial_Shareholders.number_and_class_of_shares_held", missing=missing),
            _pull(
                rec,
                sec,
                "Substantial_Shareholders.approximate_percentage_of_shareholding_of_each_class_of_shares",
                missing=missing,
            ),
            _pull(
                rec,
                sec,
                "Substantial_Shareholders.approximate_percentage_of_shareholding_in_the_issued_and_outstanding_share_capital",
                missing=missing,
            ),
        ]
    )

    ws2 = wb.create_sheet("Share capital")
    ws2.append(["Metric", "Value"])
    for label, dotted in [
        ("Authorized share capital", "Share_Capital.authorized_share_capital"),
        ("Issued share capital", "Share_Capital.issued_share_capital"),
        ("Share capital after offering", "Share_Capital.share_capital_after_offering"),
        ("Voting rights structure", "Share_Capital.voting_rights_structure"),
        ("Repurchase mandate", "Share_Capital.repurchase_mandate_details"),
    ]:
        ws2.append([label, _pull(rec, cap_sec, dotted, missing=missing)])

    out_path = out_dir / "shareholder_list.xlsx"
    out_dir.mkdir(parents=True, exist_ok=True)
    wb.save(str(out_path))
    return FillResult(
        path=out_path,
        relative_path="templated/shareholder_list.xlsx",
        schema_a_field_id="issuer.corp.shareholder_register",
        section_hint="Substantial_Shareholders",
        source_record_paths=[f"{sec}.*", f"{cap_sec}.*"],
        missing_fields=missing,
    )


# ------------------------------------------------------------------- financial_model


def _fill_financial_model(
    rec: dict[str, Any], out_dir: Path, doc_id: str
) -> FillResult:
    from openpyxl import Workbook

    missing: list[str] = []
    sec = "section_Financial_Information"

    wb = Workbook()
    header = wb.active
    header.title = "Cover"
    header.append(["Financial model (skeleton)", ""])
    header.append(["Source document", f"{doc_id}.pdf"])
    header.append(
        ["Issuer name", _pull(rec, sec, "Financial_Information.issuer_name", missing=missing)]
    )
    header.append(
        [
            "Financial years covered",
            _pull(rec, sec, "Financial_Information.financial_years", missing=missing),
        ]
    )
    header.append([])
    header.append(
        [
            "Note",
            "Cells rendered as DATA_MISSING have no supporting data in the "
            "reverse-engineered record; original audit workpapers required.",
        ]
    )

    ws = wb.create_sheet("Income statement")
    ws.append(["Metric", "Value"])
    for label, dotted in [
        ("Revenue", "Financial_Information.revenue"),
        ("Net income", "Financial_Information.net_income"),
        ("Segment data", "Financial_Information.segment_data"),
    ]:
        ws.append([label, _pull(rec, sec, dotted, missing=missing)])

    bs = wb.create_sheet("Balance sheet")
    bs.append(["Metric", "Value"])
    for label, dotted in [
        ("Total assets", "Financial_Information.assets"),
        ("Total liabilities", "Financial_Information.liabilities"),
    ]:
        bs.append([label, _pull(rec, sec, dotted, missing=missing)])

    cf = wb.create_sheet("Cash flow")
    cf.append(["Metric", "Value"])
    cf.append(
        [
            "Cash flows summary",
            _pull(rec, sec, "Financial_Information.cash_flows", missing=missing),
        ]
    )

    kpi = wb.create_sheet("KPIs")
    kpi.append(["Metric", "Value"])
    kpi.append(
        [
            "Key financial ratios",
            _pull(rec, sec, "Financial_Information.key_ratios", missing=missing),
        ]
    )

    out_path = out_dir / "financial_model.xlsx"
    out_dir.mkdir(parents=True, exist_ok=True)
    wb.save(str(out_path))
    return FillResult(
        path=out_path,
        relative_path="templated/financial_model.xlsx",
        schema_a_field_id="prof.accountants.financial_model",
        section_hint="Financial_Information",
        source_record_paths=[f"{sec}.*"],
        missing_fields=missing,
    )


# --------------------------------------------------------------- comfort_letter_draft


def _fill_comfort_letter(
    rec: dict[str, Any], out_dir: Path, doc_id: str
) -> FillResult:
    from docx import Document

    missing: list[str] = []
    cov = "section_Cover"
    fin = "section_Financial_Information"

    issuer = _pull(rec, cov, "Cover.issuer_name_en", missing=missing)
    years = _pull(rec, fin, "Financial_Information.financial_years", missing=missing)
    revenue = _pull(rec, fin, "Financial_Information.revenue", missing=missing)
    net_income = _pull(rec, fin, "Financial_Information.net_income", missing=missing)
    key_ratios = _pull(rec, fin, "Financial_Information.key_ratios", missing=missing)

    doc = Document()
    doc.add_heading("Comfort letter (draft)", level=0)
    doc.add_paragraph(
        f"To: The Sponsor.   Re: {issuer} — historical financial information "
        f"for the periods covered by this offering."
    )
    doc.add_paragraph(
        "We have performed the procedures agreed with the Sponsor in respect of "
        "the historical financial information presented in the Prospectus. Based "
        "on those procedures, and subject to the limitations described below, "
        "nothing has come to our attention that would cause us to believe that "
        "the key financial metrics extracted from the accounting records are not "
        "consistent, in all material respects, with the underlying general ledger."
    )
    doc.add_heading("Key financial reference table", level=1)
    tbl = doc.add_table(rows=4, cols=2)
    tbl.style = "Light Grid Accent 1"
    rows = [
        ("Financial years covered", years),
        ("Revenue (historical periods)", revenue),
        ("Net income (historical periods)", net_income),
        ("Key ratios", key_ratios),
    ]
    for i, (k, v) in enumerate(rows):
        tbl.rows[i].cells[0].text = k
        tbl.rows[i].cells[1].text = v
    doc.add_paragraph(
        "Limitations: this draft is reverse-engineered from the issued "
        "prospectus and is NOT a signed comfort letter; issuance requires the "
        "reporting accountant's engagement partner sign-off per AG 3.340."
    )

    out_path = out_dir / "comfort_letter_draft.docx"
    out_dir.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    return FillResult(
        path=out_path,
        relative_path="templated/comfort_letter_draft.docx",
        schema_a_field_id="prof.accountants.comfort_letter",
        section_hint="Financial_Information",
        source_record_paths=[f"{fin}.*", f"{cov}.*"],
        missing_fields=missing,
    )


# --------------------------------------------------------------------- orchestration


TEMPLATE_FILLERS = (
    _fill_issuer_corporate_info,
    _fill_shareholder_list,
    _fill_financial_model,
    _fill_comfort_letter,
)


def fill_document(
    doc_id: str,
    *,
    records_dir: Path,
    out_root: Path,
) -> list[FillResult]:
    rec_path = records_dir / f"{doc_id}.json"
    if not rec_path.exists():
        return []
    raw = json.loads(rec_path.read_text(encoding="utf-8"))
    rec = raw.get("record") or {}
    out_dir = out_root / doc_id / "templated"
    results: list[FillResult] = []
    for filler in TEMPLATE_FILLERS:
        try:
            results.append(filler(rec, out_dir, doc_id))
        except Exception as exc:  # noqa: BLE001
            print(f"  [template error] {doc_id} / {filler.__name__}: {exc}")
    return results


def _doc_ids(
    records_dir: Path,
    *,
    only: str | None,
    limit: int | None,
) -> Iterable[str]:
    if only:
        yield only
        return
    ids = sorted(p.stem for p in records_dir.glob("*.json"))
    if limit:
        ids = ids[:limit]
    yield from ids


def main() -> None:
    parser = argparse.ArgumentParser(
        description="Fill phase-1 native-doc templates from Schema B records."
    )
    parser.add_argument(
        "--records-dir", default="prospectus_kg_output/inputs/records"
    )
    parser.add_argument(
        "--out-root", default="prospectus_kg_output/native_docs"
    )
    parser.add_argument("--only", default=None, help="Fill only this doc_id")
    parser.add_argument("--limit", type=int, default=None)
    args = parser.parse_args()

    records_dir = Path(args.records_dir)
    out_root = Path(args.out_root)
    out_root.mkdir(parents=True, exist_ok=True)

    total_docs = 0
    total_files = 0
    total_missing = 0
    for doc_id in _doc_ids(records_dir, only=args.only, limit=args.limit):
        results = fill_document(doc_id, records_dir=records_dir, out_root=out_root)
        total_docs += 1
        total_files += len(results)
        doc_missing = sum(len(r.missing_fields) for r in results)
        total_missing += doc_missing
        print(
            f"[template] {doc_id}: {len(results)} file(s), {doc_missing} DATA_MISSING"
        )

    print(
        f"\nDone: {total_files} file(s) across {total_docs} doc(s); "
        f"{total_missing} DATA_MISSING field(s)."
    )


if __name__ == "__main__":
    main()
